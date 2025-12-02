"""
Data Loader for Multiple File Formats
======================================
Load raw data from CSV, Excel, Word, PDF, and text files.
"""

import os
import csv
import json
from typing import List, Dict, Optional, Union
from pathlib import Path
import config


class DataLoader:
    """Load data from various file formats."""
    
    @staticmethod
    def load(file_path: str) -> Dict:
        """Load data from any supported format.
        
        Args:
            file_path: Path to data file
            
        Returns:
            Dictionary with:
                - 'format': File format
                - 'content': Raw text content
                - 'structured': Structured data (if applicable)
                - 'metadata': File metadata
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = path.suffix.lower()
        
        if ext == '.csv':
            return DataLoader._load_csv(file_path)
        elif ext == '.xlsx':
            return DataLoader._load_excel(file_path)
        elif ext == '.docx':
            return DataLoader._load_word(file_path)
        elif ext == '.pdf':
            return DataLoader._load_pdf(file_path)
        elif ext == '.txt':
            return DataLoader._load_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    @staticmethod
    def _load_csv(file_path: str) -> Dict:
        """Load CSV file.
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Data dictionary
        """
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            # Try to detect delimiter
            sample = f.read(4096)
            f.seek(0)
            
            # Use csv.Sniffer to detect format
            try:
                dialect = csv.Sniffer().sniff(sample)
                reader = csv.DictReader(f, dialect=dialect)
            except:
                # Fallback to default
                reader = csv.DictReader(f)
            
            rows = list(reader)
        
        # Create text representation
        if rows:
            headers = list(rows[0].keys())
            content_lines = [','.join(headers)]
            for row in rows[:100]:  # Sample first 100 rows
                content_lines.append(','.join(str(row.get(h, '')) for h in headers))
            content = '\n'.join(content_lines)
        else:
            content = ""
        
        return {
            'format': 'csv',
            'content': content,
            'structured': rows,
            'metadata': {
                'rows': len(rows),
                'columns': list(rows[0].keys()) if rows else [],
                'file_name': Path(file_path).name
            }
        }
    
    @staticmethod
    def _load_excel(file_path: str) -> Dict:
        """Load Excel file.
        
        Args:
            file_path: Path to Excel file
            
        Returns:
            Data dictionary
        """
        try:
            import openpyxl
        except ImportError:
            raise ImportError("openpyxl is required for Excel support. Install with: pip install openpyxl")
        
        wb = openpyxl.load_workbook(file_path, data_only=True)
        sheet = wb.active
        
        # Extract as list of dictionaries
        rows = []
        headers = []
        
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(cell) if cell is not None else f"Column_{j}" for j, cell in enumerate(row)]
            else:
                row_dict = {headers[j]: str(cell) if cell is not None else "" 
                           for j, cell in enumerate(row) if j < len(headers)}
                rows.append(row_dict)
        
        # Create text representation
        if rows:
            content_lines = [','.join(headers)]
            for row in rows[:100]:
                content_lines.append(','.join(row.get(h, '') for h in headers))
            content = '\n'.join(content_lines)
        else:
            content = ""
        
        return {
            'format': 'excel',
            'content': content,
            'structured': rows,
            'metadata': {
                'rows': len(rows),
                'columns': headers,
                'file_name': Path(file_path).name
            }
        }
    
    @staticmethod
    def _load_word(file_path: str) -> Dict:
        """Load Word document.
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Data dictionary
        """
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx is required for Word support. Install with: pip install python-docx")
        
        doc = docx.Document(file_path)
        
        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = '\n'.join(paragraphs)
        
        # Extract tables if any
        tables_data = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                table_rows.append([cell.text for cell in row.cells])
            tables_data.append(table_rows)
        
        return {
            'format': 'word',
            'content': content,
            'structured': {
                'paragraphs': paragraphs,
                'tables': tables_data
            },
            'metadata': {
                'paragraphs': len(paragraphs),
                'tables': len(tables_data),
                'file_name': Path(file_path).name
            }
        }
    
    @staticmethod
    def _load_pdf(file_path: str) -> Dict:
        """Load PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Data dictionary
        """
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF support. Install with: pip install PyPDF2")
        
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            
            pages_text = []
            for page in reader.pages:
                pages_text.append(page.extract_text())
            
            content = '\n\n'.join(pages_text)
        
        return {
            'format': 'pdf',
            'content': content,
            'structured': {
                'pages': pages_text
            },
            'metadata': {
                'pages': len(pages_text),
                'file_name': Path(file_path).name
            }
        }
    
    @staticmethod
    def _load_text(file_path: str) -> Dict:
        """Load plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Data dictionary
        """
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        
        lines = content.split('\n')
        
        return {
            'format': 'text',
            'content': content,
            'structured': {
                'lines': lines
            },
            'metadata': {
                'lines': len(lines),
                'characters': len(content),
                'file_name': Path(file_path).name
            }
        }
    
    @staticmethod
    def get_content_sample(data: Dict, max_chars: int = 3000) -> str:
        """Get a representative sample of content for LLM analysis.
        
        Args:
            data: Data dictionary from load()
            max_chars: Maximum characters to return
            
        Returns:
            Content sample string
        """
        content = data.get('content', '')
        
        if len(content) <= max_chars:
            return content
        
        # Try to get a balanced sample
        # Take first 40%, middle 20%, last 40%
        chunk_size = max_chars // 3
        
        start = content[:chunk_size]
        middle_pos = len(content) // 2 - chunk_size // 2
        middle = content[middle_pos:middle_pos + chunk_size]
        end = content[-chunk_size:]
        
        return f"{start}\n\n...[middle content]...\n\n{middle}\n\n...[end content]...\n\n{end}"
